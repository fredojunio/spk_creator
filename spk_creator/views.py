from django.shortcuts import render
from django.http import HttpResponse
from django.conf import settings
from docx import Document
import os
import io


def index(request):
    return render(request, "index.html")


def create_spk(request):
    # Define the path to the template document
    template_path = os.path.join(
        settings.BASE_DIR, "templates", "documents", "template-SPK.docx"
    )

    # Load the template document
    template_doc = Document(template_path)

    # Assign variable
    var_year = request.POST["var_year"]
    var_loc = request.POST["var_loc"]
    var_month = request.POST["var_month"]
    var_to = request.POST["var_to"]
    var_cc = request.POST["var_cc"]
    var_type = request.POST["var_type"]
    var_branch = request.POST["var_branch"]
    var_initial = request.POST["var_initial"]
    var_router = request.POST["var_router"]
    var_address = request.POST["var_address"]
    var_before = request.POST["var_before"]
    var_after = request.POST["var_after"]
    var_pic = request.POST["var_pic"]
    var_phone = request.POST["var_phone"]
    var_service = request.POST["var_service"]
    var_by = request.POST["var_by"]

    # Define the string to find and the replacement
    old_strings = [
        "var_year",
        "var_loc",
        "var_month",
        "var_to",
        "var_cc",
        "var_type",
        "var_branch",
        "var_initial",
        "var_router",
        "var_address",
        "var_before",
        "var_after",
        "var_pic",
        "var_phone",
        "var_service",
        "var_by",
    ]
    new_strings = [
        var_year,
        var_loc,
        var_month,
        var_to,
        var_cc,
        var_type,
        var_branch,
        var_initial,
        var_router,
        var_address,
        var_before,
        var_after,
        var_pic,
        var_phone,
        var_service,
        var_by,
    ]

    # Replace text in paragraphs
    for paragraph in template_doc.paragraphs:
        for old_string, new_string in zip(old_strings, new_strings):
            for run in paragraph.runs:
                replace_text(run, old_string, new_string)

    # Replace text in tables
    for table in template_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old_string, new_string in zip(old_strings, new_strings):
                        for run in paragraph.runs:
                            replace_text(run, old_string, new_string)

    # Save the modified document to an in-memory file
    modified_doc_io = io.BytesIO()
    template_doc.save(modified_doc_io)
    modified_doc_io.seek(0)

    # Create a response with the modified document
    response = HttpResponse(
        modified_doc_io.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = (
        f'attachment; filename="SPK Upgrade Bandwidth {var_type} {var_branch}.docx"'
    )

    return response


# Function to replace text in a run
def replace_text(run, old, new):
    if old in run.text:
        run.text = run.text.replace(old, new)
